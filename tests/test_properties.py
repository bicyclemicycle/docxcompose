from datetime import datetime
from docx import Document
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docxcompose.properties import CustomProperties
from docxcompose.utils import xpath
from utils import docx_path


XPATH_CACHED_DOCPROPERTY_VALUES = 'w:r[preceding-sibling::w:r/w:fldChar/@w:fldCharType="separate"]/w:t'


class TestUpdateAllDocproperties(object):

    def test_updates_doc_properties_with_umlauts(self):
        document = Document(docx_path('outdated_docproperty_with_umlauts.docx'))

        text = xpath(
            document.element.body,
            u'.//w:fldSimple[contains(@w:instr, \'DOCPROPERTY "F\xfc\xfc"\')]//w:t')
        assert u'xxx' == text[0].text

        CustomProperties(document).update_all()

        text = xpath(
            document.element.body,
            u'.//w:fldSimple[contains(@w:instr, \'DOCPROPERTY "F\xfc\xfc"\')]//w:t')
        assert u'j\xe4ja.' == text[0].text

    def test_complex_docprop_fields_with_multiple_textnodes_are_updated(self):
        document = Document(docx_path('spellchecked_docproperty.docx'))
        paragraphs = xpath(document.element.body, '//w:p')
        assert 1 == len(paragraphs), 'input file contains one paragraph'
        assert 1 == len(xpath(document.element.body, '//w:instrText')), \
            'input contains one complex field docproperty'
        w_p = paragraphs[0]
        cached_value = xpath(w_p, XPATH_CACHED_DOCPROPERTY_VALUES)
        assert 4 == len(cached_value), \
            'doc property value is scattered over 4 parts'
        assert 'i will be spllchecked!' == ''.join(
            each.text for each in cached_value)

        CustomProperties(document).update_all()

        w_p = xpath(document.element.body, '//w:p')[0]
        cached_value = xpath(w_p, XPATH_CACHED_DOCPROPERTY_VALUES)
        assert 1 == len(cached_value), \
            'doc property value has been reset to one cached value'
        assert 'i will be spllchecked!' == cached_value[0].text

    def test_complex_docprop_with_multiple_textnode_in_same_run_are_updated(self):
        document = Document(docx_path('two_textnodes_in_run_docproperty.docx'))
        paragraphs = xpath(document.element.body, '//w:p')
        assert 1 == len(paragraphs), 'input file contains one paragraph'
        assert 1 == len(xpath(document.element.body, '//w:instrText')), \
            'input contains one complex field docproperty'

        w_p = paragraphs[0]
        cached_value = xpath(w_p, XPATH_CACHED_DOCPROPERTY_VALUES)
        assert 2 == len(cached_value), \
            'doc property value is scattered over 2 parts'
        assert 'Hello there' == ''.join(
            each.text for each in cached_value)

        CustomProperties(document).update_all()

        w_p = xpath(document.element.body, '//w:p')[0]
        cached_value = xpath(w_p, XPATH_CACHED_DOCPROPERTY_VALUES)
        assert 1 == len(cached_value), \
            'doc property value has been reset to one cached value'
        assert 'i will be spllchecked!' == cached_value[0].text

    def test_two_complex_docprop_in_same_paragraph(self):
        document = Document(docx_path('two_props_in_same_paragraph.docx'))
        assert 1 == len(document.paragraphs), 'input file should contains one paragraph'
        paragraph = document.paragraphs[0]
        assert 2 == len(xpath(paragraph._p, './/w:instrText')), \
            'input should contain two complex field docproperties'

        assert u'Foo Bar / 0' == paragraph.text

        CustomProperties(document).update_all()

        assert u'Bar / 2' == paragraph.text

    def test_multiple_identical_docprops_get_updated(self):
        document = Document(docx_path('multiple_identical_properties.docx'))
        assert 3 == len(document.paragraphs), 'input file should contain 3 paragraphs'
        for paragraph in document.paragraphs:
            assert 1 == len(xpath(paragraph._p, './/w:instrText')), \
                'paragraph should contain one complex field docproperties'

            assert u'Foo' == paragraph.text

        CustomProperties(document).update_all()

        for i, paragraph in enumerate(document.paragraphs):
            assert u'Bar' == paragraph.text, 'docprop {} was not updated'.format(i+1)


class TestUpdateSpecificDocproperty(object):

    def test_simple_field_gets_updated(self):
        document = Document(docx_path('outdated_docproperty_with_umlauts.docx'))
        text = xpath(
            document.element.body,
            u'.//w:fldSimple[contains(@w:instr, \'DOCPROPERTY "F\xfc\xfc"\')]//w:t')
        assert u'xxx' == text[0].text

        CustomProperties(document).update(u"F\xfc\xfc", u"new v\xe4lue")

        text = xpath(
            document.element.body,
            u'.//w:fldSimple[contains(@w:instr, \'DOCPROPERTY "F\xfc\xfc"\')]//w:t')
        assert u"new v\xe4lue" == text[0].text

    def test_complex_field_gets_updated(self):
        document = Document(docx_path('docproperties.docx'))
        assert 6 == len(document.paragraphs), 'input file should contain 6 paragraphs'

        properties = xpath(document.element.body, './/w:instrText')
        assert 5 == len(properties),\
            'input should contain five complex field docproperties'

        expected_paragraphs = [u'Custom Doc Properties',
                               u'Text: Foo Bar',
                               u'Number: 123',
                               u'Boolean: Y',
                               u'Date: 11.06.2019',
                               u'Float: 1.1']
        actual_paragraphs = [paragraph.text for paragraph in document.paragraphs]
        assert actual_paragraphs == expected_paragraphs

        CustomProperties(document).update("Number Property", 423)

        expected_paragraphs[2] = u'Number: 423'
        actual_paragraphs = [paragraph.text for paragraph in document.paragraphs]
        assert actual_paragraphs == expected_paragraphs

    def test_multiple_identical_docprops_get_updated(self):
        document = Document(docx_path('multiple_identical_properties.docx'))
        assert 3 == len(document.paragraphs), 'input file should contain 3 paragraphs'
        for paragraph in document.paragraphs:
            assert 1 == len(xpath(paragraph._p, './/w:instrText')), \
                'paragraph should contain one complex field docproperties'

            assert u'Foo' == paragraph.text

        CustomProperties(document).update("Text Property", "New value")

        for i, paragraph in enumerate(document.paragraphs):
            assert u'New value' == paragraph.text,\
                'docprop {} was not updated'.format(i+1)


class TestRemoveField(object):

    def test_removes_simple_field_but_keeps_value(self):
        document = Document(docx_path('outdated_docproperty_with_umlauts.docx'))
        assert 1 == len(document.paragraphs), 'input file should contain 1 paragraph'
        fields = xpath(
            document.element.body,
            u'.//w:fldSimple[contains(@w:instr, \'DOCPROPERTY "F\xfc\xfc"\')]//w:t')
        assert 1 == len(fields), 'should contain one simple field docproperty'

        assert u'Hie chund ds property: ' == document.paragraphs[0].text
        assert u'xxx' == fields[0].text

        CustomProperties(document).remove_field(u"F\xfc\xfc")
        fields = xpath(
            document.element.body,
            u'.//w:fldSimple[contains(@w:instr, \'DOCPROPERTY "F\xfc\xfc"\')]//w:t')
        assert 0 == len(fields), 'should not contain any docproperties anymore'
        # when simple field is removed, the value is moved one up in the hierarchy
        assert u'Hie chund ds property: xxx' == document.paragraphs[0].text

    def test_removes_complex_field_but_keeps_value(self):
        # test fails because field has 2 spaces before docprop name
        document = Document(docx_path('docproperties.docx'))
        assert 6 == len(document.paragraphs), 'input file should contain 6 paragraphs'

        properties = xpath(document.element.body, './/w:instrText')
        assert 5 == len(properties),\
            'input should contain five complex field docproperties'

        expected_paragraphs = [u'Custom Doc Properties',
                               u'Text: Foo Bar',
                               u'Number: 123',
                               u'Boolean: Y',
                               u'Date: 11.06.2019',
                               u'Float: 1.1']
        actual_paragraphs = [paragraph.text for paragraph in document.paragraphs]
        assert actual_paragraphs == expected_paragraphs

        CustomProperties(document).remove_field("Number Property")

        actual_paragraphs = [paragraph.text for paragraph in document.paragraphs]
        assert actual_paragraphs == expected_paragraphs

        properties = xpath(document.element.body, './/w:instrText')
        assert 4 == len(properties),\
            'only 4 fields should remain after removal of one'

    def test_removes_all_instances_of_given_field(self):
        document = Document(docx_path('multiple_identical_properties.docx'))
        assert 3 == len(document.paragraphs), 'input file should contain 3 paragraphs'
        assert 3 == len(xpath(document.element.body, './/w:instrText')), \
            'document should contain three complex field docproperties'

        for paragraph in document.paragraphs:
            assert u'Foo' == paragraph.text

        CustomProperties(document).remove_field("Text Property")

        assert 3 == len(document.paragraphs)
        assert 0 == len(xpath(document.element.body, './/w:instrText')), \
            'document should not contain any complex field anymore'
        for paragraph in document.paragraphs:
            assert u'Foo' == paragraph.text, "value should have been kept in document"

    def test_two_complex_docprop_in_same_paragraph(self):
        document = Document(docx_path('two_props_in_same_paragraph.docx'))
        assert 1 == len(document.paragraphs), 'input file should contains one paragraph'
        paragraph = document.paragraphs[0]
        assert 2 == len(xpath(paragraph._p, './/w:instrText')), \
            'input should contain two complex field docproperties'

        assert u'Foo Bar / 0' == paragraph.text

        CustomProperties(document).update_all()

        assert u'Bar / 2' == paragraph.text

    def test_removing_field_when_two_complex_docprop_in_same_paragraph(self):
        document = Document(docx_path('two_props_in_same_paragraph.docx'))
        assert 1 == len(document.paragraphs), 'input file should contains one paragraph'
        paragraph = document.paragraphs[0]
        assert 2 == len(xpath(paragraph._p, './/w:instrText')), \
            'input should contain two complex field docproperties'

        assert u'Foo Bar / 0' == paragraph.text

        CustomProperties(document).remove_field("Text Property")

        assert 1 == len(document.paragraphs)
        assert 1 == len(xpath(document.element.body, './/w:instrText')), \
            'document should contain one complex field after removal'
        assert u'Foo Bar / 0' == paragraph.text


def test_get_doc_properties():
    document = Document(docx_path('docproperties.docx'))
    props = CustomProperties(document)

    assert props['Text Property'] == 'Foo Bar'
    assert props['Number Property'] == 123
    assert props['Boolean Property'] is True
    assert props['Date Property'] == datetime(2019, 6, 11, 10, 0)

    assert props.get('Text Property') == 'Foo Bar'
    assert props.get('Number Property') == 123
    assert props.get('Boolean Property') is True
    assert props.get('Date Property') == datetime(2019, 6, 11, 10, 0)


def test_add_doc_properties():
    document = Document(docx_path('docproperties.docx'))
    props = CustomProperties(document)

    props.add('My Text Property', 'foo bar')
    assert props.get('My Text Property') == 'foo bar'

    props.add('My Boolean Property', True)
    assert props.get('My Boolean Property') is True

    props.add('My Number Property', 123)
    assert props.get('My Number Property') == 123

    props.add('My Date Property', datetime(2019, 10, 23, 15, 44, 50))
    assert props.get('My Date Property') == datetime(2019, 10, 23, 15, 44, 50)


def test_set_doc_properties():
    document = Document(docx_path('docproperties.docx'))
    props = CustomProperties(document)

    props['Text Property'] = 'baz'
    assert props['Text Property'] == 'baz'

    props['Boolean Property'] = False
    assert props['Boolean Property'] is False

    props['Number Property'] = 456
    assert props['Number Property'] == 456

    props['Date Property'] = datetime(2019, 10, 20, 12, 0)
    assert props['Date Property'] == datetime(2019, 10, 20, 12, 0)


def test_delete_doc_properties():
    document = Document(docx_path('docproperties.docx'))
    props = CustomProperties(document)

    del props['Text Property']
    del props['Number Property']

    assert 'Text Property' not in props
    assert 'Number Property' not in props

    assert xpath(props._element, u'.//cp:property/@pid') == ['2', '3', '4']


def test_set_doc_property_on_document_without_properties_creates_new_part():
    document = Document(docx_path('master.docx'))
    props = CustomProperties(document)
    props['Text Property'] = 'Foo'

    assert props.part is not None
    assert props['Text Property'] == 'Foo'

    part = document.part.package.part_related_by(RT.CUSTOM_PROPERTIES)
    assert part is not None


def test_doc_properties_keys():
    document = Document(docx_path('docproperties.docx'))
    props = CustomProperties(document)

    assert props.keys() == [
        'Text Property',
        'Number Property',
        'Boolean Property',
        'Date Property',
        'Float Property',
    ]


def test_doc_properties_values():
    document = Document(docx_path('docproperties.docx'))
    props = CustomProperties(document)

    assert props.values() == [
        'Foo Bar', 123, True, datetime(2019, 6, 11, 10, 0), 1.1]


def test_doc_properties_items():
    document = Document(docx_path('docproperties.docx'))
    props = CustomProperties(document)

    assert props.items() == [
        ('Text Property', 'Foo Bar'),
        ('Number Property', 123),
        ('Boolean Property', True),
        ('Date Property', datetime(2019, 6, 11, 10, 0)),
        ('Float Property', 1.1),
    ]
